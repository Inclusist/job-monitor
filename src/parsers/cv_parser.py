"""
CV Parser - Extract text from PDF, DOCX, and TXT files
"""

import os
import hashlib
from typing import Tuple
import chardet


class CVParser:
    """Parser for extracting text from CV files"""

    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

    @staticmethod
    def extract_text(file_path: str) -> Tuple[str, str]:
        """
        Extract text from PDF, DOCX, or TXT file

        Args:
            file_path: Path to CV file

        Returns:
            Tuple of (extracted_text, status)
            status: 'success', 'partial', or 'failed'
        """
        if not os.path.exists(file_path):
            return "", "failed"

        # Get file extension
        ext = os.path.splitext(file_path)[1].lower()

        try:
            if ext == '.pdf':
                text = CVParser.parse_pdf(file_path)
            elif ext == '.docx':
                text = CVParser.parse_docx(file_path)
            elif ext == '.txt':
                text = CVParser.parse_txt(file_path)
            else:
                return "", "failed"

            # Check if we got meaningful content
            if len(text.strip()) < 100:
                return text, "partial"

            return text, "success"

        except Exception as e:
            print(f"Error parsing {file_path}: {e}")
            return "", "failed"

    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """
        Extract text from PDF using pdfplumber

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text
        """
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber is required. Install with: pip install pdfplumber")

        text_parts = []

        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    print(f"Warning: Could not extract text from page {page_num}: {e}")
                    continue

        return "\n\n".join(text_parts)

    @staticmethod
    def parse_docx(file_path: str) -> str:
        """
        Extract text from DOCX using python-docx

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required. Install with: pip install python-docx")

        doc = Document(file_path)
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n".join(text_parts)

    @staticmethod
    def parse_txt(file_path: str) -> str:
        """
        Read plain text file with encoding detection

        Args:
            file_path: Path to text file

        Returns:
            File contents
        """
        # Detect encoding
        with open(file_path, 'rb') as f:
            raw_data = f.read()

        # Try to detect encoding
        try:
            detected = chardet.detect(raw_data)
            encoding = detected['encoding'] if detected['encoding'] else 'utf-8'
        except:
            encoding = 'utf-8'

        # Read with detected encoding
        try:
            return raw_data.decode(encoding)
        except:
            # Fallback to utf-8 with error handling
            return raw_data.decode('utf-8', errors='ignore')

    @staticmethod
    def calculate_hash(file_path: str) -> str:
        """
        Calculate SHA-256 hash of file for duplicate detection

        Args:
            file_path: Path to file

        Returns:
            SHA-256 hash as hex string
        """
        sha256_hash = hashlib.sha256()

        with open(file_path, "rb") as f:
            # Read file in chunks to handle large files
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)

        return sha256_hash.hexdigest()

    @staticmethod
    def validate_cv_file(file_path: str) -> Tuple[bool, str]:
        """
        Validate CV file before processing

        Args:
            file_path: Path to file

        Returns:
            Tuple of (is_valid, error_message)
        """
        # Check file exists
        if not os.path.exists(file_path):
            return False, "File not found"

        # Check file size
        size = os.path.getsize(file_path)
        if size > CVParser.MAX_FILE_SIZE:
            size_mb = size / (1024 * 1024)
            return False, f"File too large ({size_mb:.1f}MB). Maximum: 10MB"

        if size == 0:
            return False, "File is empty"

        # Check extension
        ext = os.path.splitext(file_path)[1].lower()
        if ext not in ['.pdf', '.docx', '.txt']:
            return False, f"Unsupported file type: {ext}. Supported: PDF, DOCX, TXT"

        return True, ""

    @staticmethod
    def save_extracted_text(text: str, output_path: str):
        """
        Save extracted text to file for caching

        Args:
            text: Extracted text
            output_path: Path to save text file
        """
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)

    @staticmethod
    def get_file_info(file_path: str) -> dict:
        """
        Get file metadata

        Args:
            file_path: Path to file

        Returns:
            Dictionary with file information
        """
        if not os.path.exists(file_path):
            return {}

        stat = os.stat(file_path)
        ext = os.path.splitext(file_path)[1].lower()

        return {
            'file_name': os.path.basename(file_path),
            'file_type': ext[1:] if ext else '',  # Remove leading dot
            'file_size': stat.st_size,
            'file_size_mb': round(stat.st_size / (1024 * 1024), 2),
            'modified_date': stat.st_mtime
        }


if __name__ == "__main__":
    # Test the parser
    print("Testing CVParser...")

    test_file = "test_cv.txt"

    # Create a test text file
    with open(test_file, 'w') as f:
        f.write("""
John Doe
Software Engineer

EXPERIENCE:
Senior Software Engineer at Tech Corp (2020-Present)
- Led team of 5 engineers
- Developed microservices architecture

SKILLS:
Python, JavaScript, SQL, Docker, Kubernetes

EDUCATION:
B.S. Computer Science, MIT, 2018
""")

    # Validate file
    is_valid, error = CVParser.validate_cv_file(test_file)
    print(f"Validation: {is_valid} - {error}")

    # Extract text
    text, status = CVParser.extract_text(test_file)
    print(f"Extraction status: {status}")
    print(f"Extracted text length: {len(text)} characters")
    print(f"Preview: {text[:200]}...")

    # Calculate hash
    file_hash = CVParser.calculate_hash(test_file)
    print(f"File hash: {file_hash}")

    # Get file info
    info = CVParser.get_file_info(test_file)
    print(f"File info: {info}")

    # Clean up
    os.remove(test_file)

    print("\nCVParser test completed successfully!")
